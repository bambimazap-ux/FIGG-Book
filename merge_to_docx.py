import os
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# File paths
CHAPTERS_DIR = r"C:\Users\nafei\Documents\מופ\סקירות ספרות\FIGG\פרקים"
SYLLABUS_PATH = r"C:\Users\nafei\Documents\מופ\סקירות ספרות\FIGG\פרקים\סילבוס_מוצע_לקורס_FIGG.md"
OUTPUT_FILE = r"C:\Users\nafei\Documents\מופ\סקירות ספרות\FIGG\ספר_קורס_FIGG.docx"

# Ordered files following the syllabus structure:
# Chapters 1-11, then 13, 14, 15, and finally Chapter 12 (Conclusion)
ORDERED_FILES = [
    "פרק_1_מבוא_לגנאלוגיה_גנטית_ו_FIGG.md",
    "פרק_2_שונות_גנטית_אנושית_ומבנה_אוכלוסיות.md",
    "פרק_3_תשתית_טכנולוגית_לנתוני_גנאלוגיה.md",
    "פרק_4_IBD_והסקת_קרבה.md",
    "פרק_5_מגבלות_חישוביות_ואוכלוסייתיות.md",
    "פרק_6_מסדי_נתונים_גנאלוגיים_וחיפוש_קרובים.md",
    "פרק_7_יישומים_אזרחיים_של_גנאלוגיה_גנטית.md",
    "פרק_8_זיהוי_נעדרים_שרידים_ו_DVI.md",
    "פרק_9_FIGG_במעבדה_הפורנזית.md",
    "פרק_10_FIGG_כתהליך_חקירתי.md",
    "פרק_11_תיאורי_מקרה_מרכזיים.md",
    "פרק_13_היבטים_משפטיים_ואתיים_של_FIGG.md",
    "פרק_14_FIGG_במשטרות_בעולם.md",
    "פרק_15_תקנים_בקרה_ואבטחת_איכות_ב_FIGG.md",
    "פרק_12_מגבלות_כשלים_ומבט_מסכם.md"
]

def set_rtl(paragraph):
    """Aligns paragraph to right and sets RTL direction in the paragraph properties."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    # Check if w:bidi exists in pPr; if not, insert at index 0 to ensure proper schema ordering
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        pPr.insert(0, OxmlElement('w:bidi'))

def set_ltr(paragraph):
    """Aligns paragraph to left and removes bidi direction (standard LTR)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is not None:
        pPr.remove(bidi)

def set_run_rtl(run):
    """Sets the run to RTL direction."""
    rPr = run.element.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rPr.insert(0, OxmlElement('w:rtl'))

def set_table_rtl(table):
    """Sets table alignment and layout flow to RTL."""
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidi')
    tblPr.append(bidi)

def add_formatted_text(paragraph, text, is_rtl=True):
    """Parses markdown bold, italic, links, and escaped brackets, writing runs to the paragraph."""
    # Clean up escaped brackets
    text = text.replace('\\[', '[').replace('\\]', ']')
    
    # Split text to parse bold and links
    pattern = re.compile(r'(\*\*.*?\*\*|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold text
        if part.startswith('**') and part.endswith('**'):
            bold_content = part[2:-2]
            # Check for links inside bold
            sub_pattern = re.compile(r'(\[.*?\]\(.*?\))')
            sub_parts = sub_pattern.split(bold_content)
            for sub_part in sub_parts:
                if sub_part.startswith('[') and ']' in sub_part and sub_part.endswith(')'):
                    match = re.match(r'\[(.*?)\]\((.*?)\)', sub_part)
                    if match:
                        link_text, link_url = match.groups()
                        run = paragraph.add_run(link_text)
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        if is_rtl:
                            set_run_rtl(run)
                            set_font_cs(run)
                else:
                    run = paragraph.add_run(sub_part)
                    run.bold = True
                    if is_rtl:
                        set_run_rtl(run)
                        set_font_cs(run)
        
        # Link text
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text, link_url = match.groups()
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
                if is_rtl:
                    set_run_rtl(run)
                    set_font_cs(run)
        
        # Regular text / check for italic
        else:
            sub_pattern_italic = re.compile(r'(\*.*?\*)')
            sub_parts_italic = sub_pattern_italic.split(part)
            for sub_part_italic in sub_parts_italic:
                if sub_part_italic.startswith('*') and sub_part_italic.endswith('*'):
                    italic_content = sub_part_italic[1:-1]
                    run = paragraph.add_run(italic_content)
                    run.italic = True
                    if is_rtl:
                        set_run_rtl(run)
                        set_font_cs(run)
                else:
                    run = paragraph.add_run(sub_part_italic)
                    if is_rtl:
                        set_run_rtl(run)
                        set_font_cs(run)

def set_font_cs(run, font_name="Arial"):
    """Ensures complex script layout uses the designated font."""
    run.font.name = font_name
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), font_name)

def add_heading_rtl(doc, text, level):
    """Creates an RTL heading paragraph and adds formatted text."""
    h = doc.add_heading(level=level)
    h.text = '' # Clear default text
    set_rtl(h)
    add_formatted_text(h, text, is_rtl=True)
    return h

def add_heading_ltr(doc, text, level):
    """Creates an LTR heading paragraph and adds formatted text."""
    h = doc.add_heading(level=level)
    h.text = '' # Clear default text
    set_ltr(h)
    add_formatted_text(h, text, is_rtl=False)
    return h

def add_heading_rtl_with_bookmark(doc, text, level, bookmark_name=None):
    """Creates a heading paragraph with RTL direction and a bookmark anchor."""
    h = doc.add_heading(level=level)
    h.text = ''
    set_rtl(h)
    
    if bookmark_name:
        # Add bookmarkStart before writing the text
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_name)
        bookmark_start.set(qn('w:name'), bookmark_name)
        h._p.append(bookmark_start)
        
    add_formatted_text(h, text, is_rtl=True)
    
    if bookmark_name:
        # Add bookmarkEnd after the text runs
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_name)
        h._p.append(bookmark_end)
        
    return h

def create_table(doc, rows):
    """Converts markdown table lines into a Word table with RTL direction."""
    parsed_rows = []
    for r in rows:
        cells = [c.strip() for c in r.split('|')]
        # Strip ends if line starts/ends with |
        if r.startswith('|'):
            cells = cells[1:]
        if r.endswith('|'):
            cells = cells[:-1]
        
        # Skip markdown separator row (e.g. |---|---|)
        is_separator = all(re.match(r'^:?-+:?$', c) for c in cells) if cells else False
        if not is_separator:
            parsed_rows.append(cells)
            
    if not parsed_rows:
        return
        
    num_cols = max(len(r) for r in parsed_rows)
    num_rows = len(parsed_rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    set_table_rtl(table)
    
    for row_idx, parsed_row in enumerate(parsed_rows):
        row = table.rows[row_idx]
        for col_idx, cell_val in enumerate(parsed_row):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                p = cell.paragraphs[0]
                set_rtl(p)
                # Bold headers
                if row_idx == 0:
                    cell_val = f"**{cell_val}**"
                add_formatted_text(p, cell_val, is_rtl=True)

def parse_references_for_chapter(file_path):
    """Parses citation details from the chapter references section."""
    ref_map = {}
    if not os.path.exists(file_path):
        return ref_map
        
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        # Matches \[1\] or [1] at start of line
        match = re.match(r'^\\?\[(\d+)\\?\]\s+(.*)$', stripped)
        if match:
            ref_idx = int(match.group(1))
            ref_val = match.group(2).strip()
            ref_map[ref_idx] = ref_val
            
    return ref_map

def extract_chapters_table_from_syllabus(syllabus_path):
    """Parses the meetings summary table from the syllabus file."""
    table_data = []
    if not os.path.exists(syllabus_path):
        return table_data
        
    with open(syllabus_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.split('\n')
    in_table = False
    for line in lines:
        stripped = line.strip()
        if '| מפגש |' in stripped:
            in_table = True
            # Header
            table_data.append(["מספר הפרק", "שם הפרק", "תיאור כללי"])
            continue
        if in_table:
            if not stripped.startswith('|'):
                in_table = False
                break
            if ':---' in stripped or '---:' in stripped:
                # Skip separator line
                continue
            # Parse row
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if len(cells) >= 3:
                meeting_num = cells[0]
                topic = cells[1]
                focus = cells[2]
                table_data.append([f"פרק {meeting_num}", topic, focus])
    return table_data

def add_introduction(doc, syllabus_path):
    """Adds a short introduction and chapters overview table at the beginning."""
    add_heading_rtl(doc, "# ספר קורס: גנאלוגיה גנטית ו־FIGG בפורנזיקה", level=1)
    
    p = doc.add_paragraph()
    set_rtl(p)
    add_formatted_text(p, "ספר קורס זה מיועד לסטודנטים בשנה השלישית לתואר ראשון במדעים, ומטרתו להקנות הבנה מדעית וחקירתית מעמיקה בתחום ה-FIGG (Forensic Investigative Genetic Genealogy). הספר מבוסס על מאמרים אקדמיים מובילים ומציג את התהליכים הביולוגיים, הטכנולוגיים והחישוביים העומדים בבסיס השיטה, לצד יישומיה המעשיים בחקירות פליליות ובזיהוי שרידים אנושיים.", is_rtl=True)
    
    add_heading_rtl(doc, "## מבנה פרקי הקורס", level=2)
    
    chapters_table_data = extract_chapters_table_from_syllabus(syllabus_path)
    
    if not chapters_table_data:
        print("Warning: Syllabus table not found. Using fallback table.")
        return
        
    table = doc.add_table(rows=len(chapters_table_data), cols=3)
    table.style = 'Table Grid'
    set_table_rtl(table)
    
    for r_idx, row_data in enumerate(chapters_table_data):
        row = table.rows[r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            set_rtl(p)
            if r_idx == 0:
                val = f"**{val}**"
            add_formatted_text(p, val, is_rtl=True)

def add_table_of_contents(doc, chapter_titles):
    """Creates a global Table of Contents linked to chapter bookmarks."""
    add_heading_rtl(doc, "## תוכן עניינים", level=2)
    
    for idx, title in enumerate(chapter_titles):
        bookmark_name = f"chapter_{idx+1}"
        p = doc.add_paragraph()
        set_rtl(p)
        
        # Word Hyperlink Element pointing to the bookmark anchor
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        hyperlink.set(qn('w:history'), '1')
        
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Style hyperlink text: Blue, underlined, and Arial CS
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), 'Arial')
        
        rPr.append(color)
        rPr.append(underline)
        rPr.append(rFonts)
        
        # Set RTL run properties inside the hyperlink
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
        
        run.append(rPr)
        
        text_elem = OxmlElement('w:t')
        text_elem.text = f"פרק {idx+1}: {title}"
        run.append(text_elem)
        
        hyperlink.append(run)
        p._p.append(hyperlink)

def get_chapter_titles(chapters_dir, ordered_files):
    """Extracts chapter titles from the first heading line of each markdown file."""
    titles = []
    for file_name in ordered_files:
        file_path = os.path.join(chapters_dir, file_name)
        if not os.path.exists(file_path):
            continue
        with open(file_path, "r", encoding="utf-8") as f:
            first_line = f.readline().strip()
            title = re.sub(r'^#+\s+', '', first_line).strip()
            title_clean = re.sub(r'^פרק\s+\d+[\s\-\:\.\,]+', '', title)
            titles.append(title_clean)
    return titles

def main():
    doc = Document()
    
    # Set default section layout to RTL
    for section in doc.sections:
        secPr = section._sectPr
        bidi = secPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            secPr.append(bidi)
            
    # Configure Normal Style for RTL & Arial
    normal_style = doc.styles['Normal']
    pPr = normal_style.element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    
    font = normal_style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Apply complex script (RTL) font properties to Normal Style
    rPr = font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Arial')
    
    # Apply complex script fonts to Heading styles
    for name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
        style = doc.styles[name]
        font = style.font
        font.name = 'Arial'
        rPr = font.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), 'Arial')
        
    print("Collecting chapter titles...")
    chapter_titles = get_chapter_titles(CHAPTERS_DIR, ORDERED_FILES)
    
    print("Writing introduction and summary table...")
    add_introduction(doc, SYLLABUS_PATH)
    
    print("Writing table of contents page...")
    doc.add_page_break()
    add_table_of_contents(doc, chapter_titles)
    
    for idx, file_name in enumerate(ORDERED_FILES):
        file_path = os.path.join(CHAPTERS_DIR, file_name)
        if not os.path.exists(file_path):
            print(f"Warning: File {file_path} not found. Skipping.")
            continue
            
        print(f"Processing {file_name} (will be mapped to Chapter {idx+1})...")
        
        # Parse chapter reference mapping for inline replacement
        ref_map = parse_references_for_chapter(file_path)
        
        # Add page break before every chapter
        doc.add_page_break()
        
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            
        lines = content.split('\n')
        i = 0
        passed_toc_header = False
        start_parsing = False
        in_references = False
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # TOC & Header stripper:
            # Skips everything up to the second occurrence of '# פרק', which follows the TOC.
            if not start_parsing:
                if stripped == '## Table of Contents':
                    passed_toc_header = True
                if passed_toc_header and stripped.startswith('# פרק'):
                    start_parsing = True
                    # Do not continue; process this line as the actual chapter heading
                else:
                    i += 1
                    continue
            
            if not stripped:
                i += 1
                continue
                
            # Table parser
            if stripped.startswith('|'):
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_rows.append(lines[i])
                    i += 1
                create_table(doc, table_rows)
                continue
                
            # Blockquote parser
            if stripped.startswith('> '):
                p = doc.add_paragraph()
                if in_references:
                    set_ltr(p)
                else:
                    set_rtl(p)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                quote_text = stripped[2:]
                add_formatted_text(p, quote_text, is_rtl=not in_references)
                i += 1
                continue
                
            # Heading parser
            h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if h_match:
                hashes, h_text = h_match.groups()
                level = len(hashes)
                
                # Check if this heading switches layout to references LTR mode
                if level == 2 and h_text.strip().lower() == 'references':
                    in_references = True
                
                # Check if this is a Further Reading citation heading (e.g. ### [1] or ### \[1\])
                ref_num_match = re.match(r'^\\?\[(\d+)\\?\]$', h_text.strip())
                if ref_num_match:
                    ref_idx = int(ref_num_match.group(1))
                    if ref_idx in ref_map:
                        # Replace the heading text with the full citation details
                        h_text = ref_map[ref_idx]
                
                # Render heading based on references section layout
                if in_references:
                    add_heading_ltr(doc, h_text, level)
                else:
                    # If Heading 1, apply TOC Bookmark and dynamically renumber chapter to idx+1
                    if level == 1 and stripped.startswith('# פרק'):
                        h_text_new = re.sub(r'^פרק\s+\d+', f'פרק {idx+1}', h_text)
                        bookmark_name = f"chapter_{idx+1}"
                        add_heading_rtl_with_bookmark(doc, h_text_new, level, bookmark_name)
                    else:
                        add_heading_rtl(doc, h_text, level)
                i += 1
                continue
                
            # Bullet list parser
            b_match = re.match(r'^(\s*)([\-\*])\s+(.*)$', line)
            if b_match:
                spaces, bullet_char, b_text = b_match.groups()
                indent_level = len(spaces) // 2
                style_name = 'List Bullet'
                if indent_level == 1:
                    style_name = 'List Bullet 2'
                elif indent_level >= 2:
                    style_name = 'List Bullet 3'
                    
                p = doc.add_paragraph(style=style_name)
                if in_references:
                    set_ltr(p)
                else:
                    set_rtl(p)
                add_formatted_text(p, b_text, is_rtl=not in_references)
                i += 1
                continue
                
            # Numbered list parser
            n_match = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
            if n_match:
                spaces, num, n_text = n_match.groups()
                indent_level = len(spaces) // 2
                style_name = 'List Number'
                if indent_level == 1:
                    style_name = 'List Number 2'
                elif indent_level >= 2:
                    style_name = 'List Number 3'
                    
                p = doc.add_paragraph(style=style_name)
                if in_references:
                    set_ltr(p)
                else:
                    set_rtl(p)
                add_formatted_text(p, n_text, is_rtl=not in_references)
                i += 1
                continue
                
            # Horizontal rule
            if stripped == '---':
                p = doc.add_paragraph()
                if in_references:
                    set_ltr(p)
                else:
                    set_rtl(p)
                p.add_run('__________________________________________________')
                i += 1
                continue
                
            # Normal paragraph
            p = doc.add_paragraph()
            if in_references:
                set_ltr(p)
            else:
                set_rtl(p)
            add_formatted_text(p, stripped, is_rtl=not in_references)
            i += 1
            
    doc.save(OUTPUT_FILE)
    print(f"Saved merged document to {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
