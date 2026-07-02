import os
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# File paths
SUMMARY_DIR = r"C:\Users\nafei\Documents\מופ\סקירות ספרות\FIGG\סיכום קורס קצר"
INPUT_MD = os.path.join(SUMMARY_DIR, "סיכום_מקצועי_בנושא_FIGG.md")
OUTPUT_DOCX = os.path.join(SUMMARY_DIR, "סיכום_מקצועי_בנושא_FIGG.docx")

def set_rtl(paragraph):
    """Aligns paragraph to right and sets RTL direction in the paragraph properties."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    # Check if w:bidi exists in pPr; if not, insert at index 0 to ensure proper schema ordering
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        pPr.insert(0, OxmlElement('w:bidi'))

def set_ltr(paragraph):
    """Aligns paragraph to left and removes bidi direction (standard LTR)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is not None:
        pPr.remove(bidi)

def set_run_rtl(run):
    """Sets the run to RTL direction."""
    rPr = run.element.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rPr.insert(0, OxmlElement('w:rtl'))

def add_formatted_text(paragraph, text, is_rtl=True):
    """Parses markdown bold, italic, links, and escaped brackets, writing runs to the paragraph."""
    # Clean up escaped brackets
    text = text.replace('\\[', '[').replace('\\]', ']')
    
    # Split text to parse bold and links
    pattern = re.compile(r'(\*\*.*?\*\*|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold text
        if part.startswith('**') and part.endswith('**'):
            bold_content = part[2:-2]
            # Check for links inside bold
            sub_pattern = re.compile(r'(\[.*?\]\(.*?\))')
            sub_parts = sub_pattern.split(bold_content)
            for sub_part in sub_parts:
                if sub_part.startswith('[') and ']' in sub_part and sub_part.endswith(')'):
                    match = re.match(r'\[(.*?)\]\((.*?)\)', sub_part)
                    if match:
                        link_text, link_url = match.groups()
                        run = paragraph.add_run(link_text)
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        if is_rtl:
                            set_run_rtl(run)
                            set_font_cs(run)
                else:
                    run = paragraph.add_run(sub_part)
                    run.bold = True
                    if is_rtl:
                        set_run_rtl(run)
                        set_font_cs(run)
        
        # Link text
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text, link_url = match.groups()
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
                if is_rtl:
                    set_run_rtl(run)
                    set_font_cs(run)
        
        # Regular text / check for italic
        else:
            sub_pattern_italic = re.compile(r'(\*.*?\*)')
            sub_parts_italic = sub_pattern_italic.split(part)
            for sub_part_italic in sub_parts_italic:
                if sub_part_italic.startswith('*') and sub_part_italic.endswith('*'):
                    italic_content = sub_part_italic[1:-1]
                    run = paragraph.add_run(italic_content)
                    run.italic = True
                    if is_rtl:
                        set_run_rtl(run)
                        set_font_cs(run)
                else:
                    run = paragraph.add_run(sub_part_italic)
                    if is_rtl:
                        set_run_rtl(run)
                        set_font_cs(run)

def set_font_cs(run, font_name="Arial"):
    """Ensures complex script layout uses the designated font."""
    run.font.name = font_name
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), font_name)

def add_heading_rtl(doc, text, level):
    """Creates an RTL heading paragraph and adds formatted text."""
    h = doc.add_heading(level=level)
    h.text = '' # Clear default text
    set_rtl(h)
    add_formatted_text(h, text, is_rtl=True)
    # Apply paragraph spacing
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_heading_ltr(doc, text, level):
    """Creates an LTR heading paragraph and adds formatted text."""
    h = doc.add_heading(level=level)
    h.text = '' # Clear default text
    set_ltr(h)
    add_formatted_text(h, text, is_rtl=False)
    # Apply paragraph spacing
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_heading_rtl_with_bookmark(doc, text, level, bookmark_name=None):
    """Creates a heading paragraph with RTL direction and a bookmark anchor."""
    h = doc.add_heading(level=level)
    h.text = ''
    set_rtl(h)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    
    if bookmark_name:
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_name)
        bookmark_start.set(qn('w:name'), bookmark_name)
        h._p.append(bookmark_start)
        
    add_formatted_text(h, text, is_rtl=True)
    
    if bookmark_name:
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_name)
        h._p.append(bookmark_end)
        
    return h

def add_table_of_contents(doc, toc_entries):
    """Creates a global Table of Contents linked to section bookmarks."""
    add_heading_rtl(doc, "## תוכן עניינים", level=2)
    
    for idx, (title, bookmark_name) in enumerate(toc_entries):
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        # Word Hyperlink Element pointing to the bookmark anchor
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        hyperlink.set(qn('w:history'), '1')
        
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Style hyperlink text: Blue, underlined, and Arial CS
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), 'Arial')
        
        rPr.append(color)
        rPr.append(underline)
        rPr.append(rFonts)
        
        # Set RTL run properties inside the hyperlink
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
        
        run.append(rPr)
        
        text_elem = OxmlElement('w:t')
        text_elem.text = title
        run.append(text_elem)
        
        hyperlink.append(run)
        p._p.append(hyperlink)

def main():
    if not os.path.exists(INPUT_MD):
        print(f"Error: input file {INPUT_MD} not found.")
        return

    print("Reading markdown file...")
    with open(INPUT_MD, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split('\n')
    
    # First pass: find all Level 2 headings in the main body (after line 30)
    # We find the second title "# סיכום מקצועי בנושא FIGG" to start counting
    title_count = 0
    toc_entries = []
    
    for line in lines:
        stripped = line.strip()
        if stripped == "# סיכום מקצועי בנושא FIGG":
            title_count += 1
            continue
        
        # Only parse headings in the main content (after the second title occurrence)
        if title_count >= 2:
            h_match = re.match(r'^##\s+(.*)$', stripped)
            if h_match:
                h_text = h_match.group(1).strip()
                # Clean up markdown styling inside headings if any
                h_text_clean = re.sub(r'[\*\`\[\]\(\)]', '', h_text)
                # Generate unique bookmark name
                bookmark_name = f"sec_{len(toc_entries) + 1}"
                toc_entries.append((h_text_clean, bookmark_name))

    print(f"Found {len(toc_entries)} main sections for Table of Contents.")

    # Initialize document
    doc = Document()
    
    # Set default margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Set section layout direction to RTL
        secPr = section._sectPr
        bidi = secPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            secPr.append(bidi)

    # Configure Normal Style for RTL & Arial
    normal_style = doc.styles['Normal']
    pPr = normal_style.element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    
    font = normal_style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    rPr = font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Arial')
    
    # Apply complex script fonts to Heading styles
    for name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
        style = doc.styles[name]
        font = style.font
        font.name = 'Arial'
        rPr = font.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), 'Arial')

    # Page 1: Main Title and Interactive TOC
    print("Writing Title page and interactive TOC...")
    add_heading_rtl(doc, "סיכום מקצועי בנושא FIGG", level=1)
    
    add_table_of_contents(doc, toc_entries)
    
    # Page Break after TOC page
    doc.add_page_break()

    # Second pass: write document body starting from the second title occurrence
    print("Writing document body...")
    title_count = 0
    in_references = False
    toc_idx = 0
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if stripped == "# סיכום מקצועי בנושא FIGG":
            title_count += 1
            # We skip the first title on page 2 since we already have page 1's title,
            # but wait, let's keep the second title as a Heading 1 on page 2 if it starts the body.
            # Actually, to preserve 100% of content and follow rules:
            # Let's write the second Title as Heading 1 on Page 2, just like in the original markdown file.
            if title_count == 2:
                # Add Heading 1 at start of page 2
                add_heading_rtl(doc, "סיכום מקצועי בנושא FIGG", level=1)
                i += 1
                continue
            else:
                i += 1
                continue
                
        if title_count < 2:
            # Skip everything before the second title (Undermind logo, static TOC, etc.)
            i += 1
            continue
            
        if not stripped:
            i += 1
            continue
            
        # Heading parser
        h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h_match:
            hashes, h_text = h_match.groups()
            level = len(hashes)
            
            # Check if this heading switches layout to references LTR mode
            if level == 2 and h_text.strip().lower() == 'references':
                in_references = True
                
            if in_references:
                add_heading_ltr(doc, h_text, level)
            else:
                if level == 2 and toc_idx < len(toc_entries):
                    # Align with the TOC bookmarks
                    _, bookmark_name = toc_entries[toc_idx]
                    add_heading_rtl_with_bookmark(doc, h_text, level, bookmark_name)
                    toc_idx += 1
                else:
                    add_heading_rtl(doc, h_text, level)
            i += 1
            continue
            
        # Bullet list parser
        b_match = re.match(r'^(\s*)([\-\*])\s+(.*)$', line)
        if b_match:
            spaces, bullet_char, b_text = b_match.groups()
            indent_level = len(spaces) // 2
            style_name = 'List Bullet'
            if indent_level == 1:
                style_name = 'List Bullet 2'
            elif indent_level >= 2:
                style_name = 'List Bullet 3'
                
            p = doc.add_paragraph(style=style_name)
            if in_references:
                set_ltr(p)
            else:
                set_rtl(p)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, b_text, is_rtl=not in_references)
            i += 1
            continue
            
        # Numbered list parser
        n_match = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if n_match:
            spaces, num, n_text = n_match.groups()
            indent_level = len(spaces) // 2
            style_name = 'List Number'
            if indent_level == 1:
                style_name = 'List Number 2'
            elif indent_level >= 2:
                style_name = 'List Number 3'
                
            p = doc.add_paragraph(style=style_name)
            if in_references:
                set_ltr(p)
            else:
                set_rtl(p)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, n_text, is_rtl=not in_references)
            i += 1
            continue
            
        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            if in_references:
                set_ltr(p)
            else:
                set_rtl(p)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add thin divider
            p.add_run('__________________________________________________')
            i += 1
            continue
            
        # Normal paragraph
        p = doc.add_paragraph()
        if in_references:
            set_ltr(p)
        else:
            set_rtl(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped, is_rtl=not in_references)
        i += 1
        
    doc.save(OUTPUT_DOCX)
    print(f"Success! Saved merged summary to {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
