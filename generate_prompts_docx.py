import os
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# File paths
WORKSPACE_DIR = r"C:\Users\nafei\Documents\מופ\סקירות ספרות\FIGG"
INPUT_MD = os.path.join(WORKSPACE_DIR, "פרומפטים_NotebookLM.md")
OUTPUT_DOCX = os.path.join(WORKSPACE_DIR, "פרומפטים_NotebookLM.docx")

def set_rtl(paragraph):
    """Aligns paragraph to right and sets RTL direction in the paragraph properties."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        pPr.insert(0, OxmlElement('w:bidi'))

def set_ltr(paragraph):
    """Aligns paragraph to left and removes bidi direction (standard LTR)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is not None:
        pPr.remove(bidi)

def set_run_rtl(run):
    """Sets the run to RTL direction."""
    rPr = run.element.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rPr.insert(0, OxmlElement('w:rtl'))

def set_font_cs(run, font_name="Arial"):
    """Ensures complex script layout uses the designated font."""
    run.font.name = font_name
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), font_name)

def add_formatted_text(paragraph, text, is_rtl=True):
    """Parses markdown bold, italic, and links, writing runs to the paragraph."""
    # Split text to parse bold and links
    pattern = re.compile(r'(\*\*.*?\*\*|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold text
        if part.startswith('**') and part.endswith('**'):
            bold_content = part[2:-2]
            run = paragraph.add_run(bold_content)
            run.bold = True
            if is_rtl:
                set_run_rtl(run)
                set_font_cs(run)
        
        # Link text
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text, link_url = match.groups()
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
                if is_rtl:
                    set_run_rtl(run)
                    set_font_cs(run)
        
        # Regular text
        else:
            run = paragraph.add_run(part)
            if is_rtl:
                set_run_rtl(run)
                set_font_cs(run)

def main():
    if not os.path.exists(INPUT_MD):
        print(f"Error: {INPUT_MD} not found.")
        return

    print("Reading markdown prompts guide...")
    with open(INPUT_MD, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Set section layout direction to RTL
        secPr = section._sectPr
        bidi = secPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            secPr.append(bidi)

    # Configure Normal Style for RTL & Arial
    normal_style = doc.styles['Normal']
    pPr = normal_style.element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    
    font = normal_style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    rPr = font.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Arial')
    
    # Configure Heading styles
    for name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        style = doc.styles[name]
        font = style.font
        font.name = 'Arial'
        rPr = font.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), 'Arial')

    print("Parsing markdown content to Word...")
    
    in_code_block = False
    
    for line in lines:
        stripped = line.strip()
        
        # Check for code block markers
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            # Add LTR prompt paragraph
            p = doc.add_paragraph()
            set_ltr(p)
            # Indent code/prompt block to make it look distinct
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            
            # Use Consolas for code blocks to make it look like a prompt box
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            # Dark gray text color
            run.font.color.rgb = RGBColor(60, 60, 60)
            continue
            
        if not stripped:
            continue
            
        # Heading parser
        h_match = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if h_match:
            hashes, h_text = h_match.groups()
            level = len(hashes)
            
            h = doc.add_heading(level=level)
            h.text = ''
            set_rtl(h)
            add_formatted_text(h, h_text, is_rtl=True)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            continue
            
        # Bullet list parser
        b_match = re.match(r'^(\s*)([\-\*])\s+(.*)$', line)
        if b_match:
            spaces, bullet_char, b_text = b_match.groups()
            indent_level = len(spaces) // 2
            style_name = 'List Bullet'
            if indent_level == 1:
                style_name = 'List Bullet 2'
            elif indent_level >= 2:
                style_name = 'List Bullet 3'
                
            p = doc.add_paragraph(style=style_name)
            set_rtl(p)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, b_text, is_rtl=True)
            continue
            
        # Normal paragraph
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped, is_rtl=True)
        
    doc.save(OUTPUT_DOCX)
    print(f"Success! Saved Word prompts document to {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
